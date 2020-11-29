from django.conf import settings
from .doc_formatter import DocumentFormatter
from docx.oxml.xmlchemy import OxmlElement
from .ImageExtractor import ImageExtractor
from docx.oxml.ns import qn
from docx.enum.style import WD_BUILTIN_STYLE

w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def prepare_to_format(path_to_document, hyphenation=False):
    '''
        Copies content from original document and makes a formatter from the copy.
        Returns a formatter with document ready to be formatted.
    '''
    
    ImageExtractor(path_to_document).extract_images()
    formatter = DocumentFormatter(path_to_document)
    formatter.remove_hyperlinks()
    headers, _ = formatter.get_headers()
    headers_text = list(map(lambda x: x.text, headers))
    if hyphenation:
        new_formatter = DocumentFormatter(settings.BASE_DIR / 'format/services/base-hyphen.docx')
    else:
        new_formatter = DocumentFormatter(settings.BASE_DIR / 'format/services/base.docx')
    new_formatter.document._body.clear_content()
    for p in formatter.document.paragraphs:
        if p.text in headers_text and p.text != headers_text[0]:
            paragraph = new_formatter.document.add_paragraph(p.text)
            new_formatter.add_page_break_before_paragraph(paragraph)
        else:
            numPr = p._p.get_or_add_pPr().numPr
            if numPr is not None:
                list_type = get_list_type_by_paragraph(formatter, p)     
                if list_type == 'bullet':
                    list_style = new_formatter.document.styles['List Bullet']
                else:
                    list_style = new_formatter.document.styles['List Number']
                copied_list = new_formatter.document.add_paragraph(style=list_style)
                copy_runs(p, copied_list)
            elif '<a:graphicData' in p._p.xml:
                print('i')
                copy_images(formatter, new_formatter, p)
            else:
                copied_paragraph = new_formatter.document.add_paragraph()
                copy_runs(p, copied_paragraph)
    return new_formatter

def copy_images(src_formatter, dest_formatter, paragraph):
    rels = src_formatter.rels
    for rId in rels:
        if rId in paragraph._p.xml:
            dest_formatter.document.add_picture(f'{settings.MEDIA_ROOT}/images/{rels[rId]}')

def copy_runs(src_paragraph, dest_paragraph):
    for run in src_paragraph.runs:
        copied_run = dest_paragraph.add_run(run.text)
        copied_run.font.bold = run.font.bold
        copied_run.font.italic = run.font.italic
        copied_run.font.underline = run.font.underline
        copied_run.font.color.rgb = run.font.color.rgb

def get_list_type_by_paragraph(formatter, paragraph):
    numPr = paragraph._p.get_or_add_pPr().numPr
    numId = numPr.numId.val
    ilvl = numPr.ilvl.val
    numbering_object = formatter.get_numbering_object()
    num_object = get_num_object_by_id(numbering_object, numId)
    abstract_num_id = get_abstract_num_id_from_num_object(num_object)
    abstract_num_object = get_abstract_num_object_by_id(numbering_object, abstract_num_id)
    levels = get_levels_from_abstract_num_object(abstract_num_object)
    level = get_level_by_ilvl(levels, ilvl)
    list_type = get_list_type_from_level(level)
    return(list_type)

def get_num_object_by_id(numbering_object, numId):
    num_object = list(filter(
        lambda x: int(x.values()[0]) == numId, numbering_object.xpath('w:num')
    ))[0]
    return num_object

def get_abstract_num_id_from_num_object(num_object):
    abstract_num_id = num_object.xpath('w:abstractNumId')[0].get(f'{{{w}}}val')
    return abstract_num_id

def get_abstract_num_object_by_id(numbering_object, abstract_num_id):
    abstract_num_object = list(filter(
        lambda x: abstract_num_id == x.get(f'{{{w}}}abstractNumId'), numbering_object.xpath('w:abstractNum')
    ))[0]
    return abstract_num_object

def get_levels_from_abstract_num_object(abstract_num_object):
    levels = abstract_num_object.xpath('w:lvl', namespaces={'w': w})
    return levels

def get_level_by_ilvl(levels, ilvl):
    level = list(filter(
        lambda x: int(x.get(f'{{{w}}}ilvl')) == ilvl, levels
    ))[0]
    return level

def get_list_type_from_level(level):
    list_type = level.xpath('w:numFmt', namespaces={'w': w})[0].get(f'{{{w}}}val')
    return list_type

def get_level_restart(level):
    level_restart = level.xpath('w:lvlRestart', namespaces={'w': w})
    return level_restart