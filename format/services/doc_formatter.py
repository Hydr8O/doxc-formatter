from docx import Document
from docx.shared import Mm, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE, WD_BUILTIN_STYLE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.parts.image import ImagePart
from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement
from .constants import HeaderConstants, TextConstants, PageConstants, ListConstants, StyleNameConstants, ImageConstants, ImageCaptionConstants


class DocumentFormatter:
    '''
        MS Word document formatter.
        Reads document in specified path and formats it.
    '''
    def __init__(self, document=None):
        '''
            Distinguish style is needed in order to emphasize certain parts of the document.
            It will be removed befor saving.
        '''
        self.document = Document(document)
        self.path = document
        self.rels = self.__map_rels_to_images()


    def __map_rels_to_images(self):
        rels = {}
        for r in self.document.part.rels.values():
            if isinstance(r._target, ImagePart):
                rels[r.rId] = r._target.partname.split('/')[-1]
        return rels

    def set_margins(self, margins):
        '''
            Sets margins to the whole document.
            margins - tuple (top, left, bottom, right).
        '''
        top_margin, left_margin, bottom_margin, right_margin = margins
        for section in self.document.sections:
            section.top_margin = top_margin
            section.left_margin = left_margin
            section.bottom_margin = bottom_margin
            section.right_margin = right_margin

    def add_list_style(self, style_name):
        '''
            Adds a list style to the document.
            Returns added style.
        '''
        added_style = self.document.styles.add_style(
            style_name,
            WD_STYLE_TYPE.LIST
        )
        return added_style

    def add_character_style(self, style_name):
        '''
            Adds a character style to the document.
            Returns added style.
        '''
        added_style = self.document.styles.add_style(
            style_name,
            WD_STYLE_TYPE.CHARACTER
        )
        return added_style

    def add_paragraph_style(self, style_name):
        '''
            Adds a paragraph style to the document.
            Returns added style.
        '''
        added_style = self.document.styles.add_style(
            style_name, 
            WD_STYLE_TYPE.PARAGRAPH
        )
        return added_style

    def insert_paragraph_after(self, paragraph, text=None, style=None):
        '''Insert a new paragraph after the given paragraph.'''
        new_p = OxmlElement('w:p')
        paragraph._p.addnext(new_p)
        new_para = Paragraph(new_p, paragraph._parent)
        if text:
            new_para.add_run(text)
        if style is not None:
            new_para.style = style
        return new_para

    def get_headers(self):
        headers = [self.document.paragraphs[0]]
        secondary_headers = []
        for i in range(1, len(self.document.paragraphs)):
            for run in self.document.paragraphs[i].runs:
                if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                    headers.append(self.document.paragraphs[i + 1])
                    break
                else:
                    if not run.bold:
                        break
                secondary_headers.append(self.document.paragraphs[i])

        return headers, secondary_headers

    def get_lists(self):
        bullet_lists = []
        number_lists = []
        for paragraph in self.document.paragraphs:
            if paragraph.style.name == 'List Number':
                number_lists.append(paragraph)
            elif paragraph.style.name == 'List Bullet':
                bullet_lists.append(paragraph)
        return {
            'bullet_lists': bullet_lists,
            'number_lists': number_lists
        }

    def get_images_and_captions(self):
        images = []
        captions = []
        for i, paragraph in enumerate(self.document.paragraphs):
            if '<a:graphicData' in paragraph._p.xml:
                images.append(paragraph)
                try:
                    captions.append(self.document.paragraphs[i + 1])
                except:
                    print('last image doesn\'t have a caption')
        return images, captions


    def show_style_in_ui(self, style):
        ''' 
            Make a style visible in MS Word's ui
        '''
        style.hidden = False
        style.quick_style = True

    def add_standard_header_style(self, style_name):
        '''
            Adds header style to the document and returns it.
            Header takes its style from values, defined in constants
        '''
        try:
            header_style = self.add_paragraph_style(style_name)
        except ValueError:
            header_style = self.document.styles[style_name]
        
        header_style.base_style = self.document.styles['Heading 1']
        header_style.font.size = HeaderConstants.FONT_SIZE
        header_style.font.name = HeaderConstants.FONT_NAME
        header_style.font.color.rgb = HeaderConstants.FONT_COLOR
        header_style.paragraph_format.alignment = HeaderConstants.ALIGNMENT
        header_style.font.all_caps = True
        header_style.paragraph_format.space_after = HeaderConstants.SPACING_AFTER
        header_style.paragraph_format.line_spacing = HeaderConstants.LINE_SPACING
        header_style.paragraph_format.space_before = HeaderConstants.SPACING_BEFORE
        return header_style

    def add_standard_text_style(self, style_name):
        '''
            Adds text style to the document and returns it.
            Text takes its style from values, defined in constants
        '''
        try:
            text_style = self.add_paragraph_style(style_name)
        except ValueError:
            text_style = self.document.styles[style_name]
        
        text_style.base_style = self.document.styles['Normal']
        text_style.font.size = TextConstants.FONT_SIZE
        text_style.font.name = TextConstants.FONT_NAME
        text_style.font.color.rgb = TextConstants.FONT_COLOR
        text_style.paragraph_format.alignment = TextConstants.ALIGNMENT
        text_style.paragraph_format.line_spacing = TextConstants.LINE_SPACING
        text_style.paragraph_format.first_line_indent = TextConstants.FIRST_LINE_INDENT
        text_style.paragraph_format.space_after = TextConstants.SPACING_AFTER
        return text_style

    def add_standard_image_style(self, style_name):
        try:
            image_style = self.add_paragraph_style(style_name)
        except ValueError:
            image_style = self.document.styles[style_name]

        image_style.paragraph_format.alignment = ImageCaptionConstants.ALIGNMENT
        image_style.paragraph_format.first_line_indent = ImageCaptionConstants.FIRST_LINE_INDENT
        

        return image_style

    def add_standard_image_caption_style(self, style_name):
        try:
            image_caption_style = self.add_paragraph_style(style_name)
        except ValueError:
            image_caption_style = self.document.styles[style_name]

        image_caption_style.paragraph_format.alignment = ImageCaptionConstants.ALIGNMENT
        image_caption_style.paragraph_format.first_line_indent = ImageCaptionConstants.FIRST_LINE_INDENT
        image_caption_style.font.size = ImageCaptionConstants.FONT_SIZE
        image_caption_style.font.name = ImageCaptionConstants.FONT_NAME
        image_caption_style.font.color.rgb = ImageCaptionConstants.FONT_COLOR
        image_caption_style.paragraph_format.space_after = ImageCaptionConstants.SPACING_AFTER
        return image_caption_style

    def add_standard_list_style(self, style_name, type):
        try:
            list_style = self.add_paragraph_style(style_name)
        except ValueError:
            list_style = self.document.styles[style_name]
        '''
            Here put list style definitions from ListsConstants
        '''
        if type == 'number':
            list_style.base_style = self.document.styles['List Number']
        elif type == 'bullet':
            list_style.base_style = self.document.styles['List Bullet']

        list_style.font.size = ListConstants.FONT_SIZE
        list_style.font.name = ListConstants.FONT_NAME
        list_style.font.color.rgb = ListConstants.FONT_COLOR
        list_style.paragraph_format.alignment = ListConstants.ALIGNMENT
        list_style.paragraph_format.line_spacing = ListConstants.LINE_SPACING
        list_style.paragraph_format.first_line_indent = ListConstants.FIRST_LINE_INDENT
        list_style.paragraph_format.space_after = ListConstants.SPACING_AFTER
        list_style.paragraph_format.left_indent = ListConstants.LEFT_INDENT
        return list_style

    def add_standard_character_style(self, style_name):
        try:
            character_style = self.add_character_style(style_name)
        except ValueError:
            character_style = self.document.styles[style_name]
        character_style.font.name = 'Times New Roman'
        character_style.font.size = Pt(14)
        '''
            Here put character style definitions from ListsConstants
        '''
        return character_style

    def style_document(
            self, 
            header_style=None, 
            text_style=None,
            bullet_style=None,
            number_style=None,
            image_style=None,
            image_caption_style=None,
            margins=None
        ):
        '''
            Styles whole document.
        '''
        if text_style is None:
            text_style = self.add_standard_text_style(StyleNameConstants.TEXT_STYLE)
            
        if bullet_style is None:
            bullet_style = self.add_standard_list_style(
                StyleNameConstants.BULLET_LIST_STYLE,
                'bullet'
            )
            
        if number_style is None:
            number_style = self.add_standard_list_style(
                StyleNameConstants.NUMBER_LIST_STYLE,
                'number'
            )
            
        if header_style is None:
            header_style = self.add_standard_header_style(StyleNameConstants.HEADER_STYLE)
            
        if image_style is None:
            image_style = self.add_standard_image_style(StyleNameConstants.IMAGE_STYLE)

        if image_caption_style is None:
            image_caption_style = self.add_standard_image_caption_style(StyleNameConstants.IMAGE_CAPTION_STYLE)

        if margins is None:
            margins = (
                PageConstants.MARGIN_TOP, 
                PageConstants.MARGIN_LEFT, 
                PageConstants.MARGIN_BOTTOM, 
                PageConstants.MARGIN_RIGHT
            )

        self.style_bullet_lists(bullet_style)
        self.style_number_lists(number_style)
        self.style_headers(header_style)
        self.style_images(image_style)
        self.style_image_captions(image_caption_style)
        self.style_text(text_style)
        self.set_margins(margins)
        

    def style_images(self, style):
        images, _ = self.get_images_and_captions()
        for image in images:
            image.style = style

    def remove_style_from_ui(self, style):
        '''
            Make a style not visible in MS Word's ui.
        '''
        style.hidden = True
        style.quick_style = False

    def add_page_break_after_paragraph(self, paragraph, text=None):
        '''
            Adds page break after paragraph and returns paragraph on the new page with text.
        '''
        next_paragraph = self.insert_paragraph_after(paragraph)
        new_header = next_paragraph.add_run()
        new_header.add_break(WD_BREAK.PAGE)
        return self.insert_paragraph_after(next_paragraph, text)

    def add_page_break_before_paragraph(self, paragraph):
        previous_paragraph = paragraph.insert_paragraph_before()
        new_header = previous_paragraph.add_run()
        new_header.add_break(WD_BREAK.PAGE)
        

    def style_headers(self, style):
        headers, _ = self.get_headers()
        for header in headers:
            header.style = style

    def style_lists(self, bullet_style=None, number_style=None):
        if bullet_style is None:
            bullet_style = self.document.styles['List Bullet']
        if number_style is None:
            number_style = self.document.styles['List Number']
        lists = self.get_lists()
        
        self.style_bullet_lists(bullet_style)
        self.style_number_lists(number_style)

    def style_bullet_lists(self, bullet_style=None):
        bullet_lists = self.get_lists()['bullet_lists']

        if bullet_style is None:
            bullet_style = self.document.styles['List Bullet']
        
        for list in bullet_lists:
            list.style = bullet_style
        
    def style_number_lists(self, number_style=None):
        number_lists = self.get_lists()['number_lists']
        if number_lists is None:
            return

        if number_style is None:
            number_style = self.document.styles['List Number']
        
        for list in number_lists:
            list.style = number_style

    def style_image_captions(self, style):
        _, captions = self.get_images_and_captions()
        
        for i, caption in enumerate(captions):
            caption.text = f'Рисунок {i + 1} — ' + caption.text
            caption.style = style

    def add_style(self, style, type):
        self.document.styles.add_style(style, type)

    def style_all_text(self, style):
        '''
            Applies a style to all text in the document
        '''
        for paragraph in self.document.paragraphs:
            paragraph.style = style

    def style_text(self, style):
        '''
            CALL ONLY AFTER STYLING THE REST OF DOCUMENT CONTENTS
            Applies a style to all paragraphs 
            except headers(first paragraph after page break)
            and lists
        '''

        for paragraph in self.document.paragraphs:
            if (
                paragraph.style.name != StyleNameConstants.HEADER_STYLE and
                paragraph.style.name != StyleNameConstants.NUMBER_LIST_STYLE and
                paragraph.style.name != StyleNameConstants.BULLET_LIST_STYLE and
                paragraph.style.name != StyleNameConstants.IMAGE_STYLE and
                paragraph.style.name != StyleNameConstants.IMAGE_CAPTION_STYLE
            ):
                paragraph.style = style
        
    def add_page_numbers(self):
        footer_paragraph = self.document.sections[0].footer.paragraphs[0]
        run = footer_paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        self._create_attribute(fldChar1, 'w:fldCharType', 'begin')

        instrText = OxmlElement('w:instrText')
        self._create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        self._create_attribute(fldChar2, 'w:fldCharType', 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

        footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def _create_attribute(self, element, name, value):
        element.set(qn(name), value)

    def remove_hyperlinks(self):
        for paragraph in self.document.paragraphs:
            self.remove_hyperlinks_from_paragraph(paragraph)

    def remove_hyperlinks_from_paragraph(self, paragraph):
        p = paragraph._p
        for link in p.xpath("./w:hyperlink"):
            link_content = link.getchildren()
            for element in link_content:
                link.addprevious(element)
            p.remove(link)
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)

    def get_numbering_object(self):
        return self.document.part.numbering_part.numbering_definitions._numbering

    def clear_list_formatting(self):
        lists = self.get_lists()
        for list in lists:
            runs = list.runs
            list.clear()
            for run in runs:
                run = list.add_run(run.text)

    def save(self, name):
        '''
            Saves formatted document using provided name.
        '''
        self.document.save(name)

#Code for making Table of Contents

# paragraph = document.add_paragraph()
# run = paragraph.add_run()
# fldChar = OxmlElement('w:fldChar')  # creates a new element
# fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
# instrText = OxmlElement('w:instrText')
# instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
# instrText.text = 'TOC \\o "1-3" \\h \\z \\u'   # change 1-3 depending on heading levels you need

# fldChar2 = OxmlElement('w:fldChar')
# fldChar2.set(qn('w:fldCharType'), 'separate')
# fldChar3 = OxmlElement('w:t')
# fldChar3.text = "Right-click to update field."
# fldChar2.append(fldChar3)

# fldChar4 = OxmlElement('w:fldChar')
# fldChar4.set(qn('w:fldCharType'), 'end')

# r_element = run._r
# r_element.append(fldChar)
# r_element.append(instrText)
# r_element.append(fldChar2)
# r_element.append(fldChar4)
# document.save('test.docx')
